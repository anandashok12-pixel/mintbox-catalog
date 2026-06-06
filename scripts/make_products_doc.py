"""Generate a Word document listing MintBox products (name, category, price).

Source: scripts/seed.ts seed catalog (repo-committed data).
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# (name, category, price INR) grouped in catalog order
PRODUCTS = [
    ("Premium Kraft Gift Box", "Packaging", 120),
    ("Luxury Rigid Gift Box", "Packaging", 280),
    ("Executive Leather Diary", "Diaries & Notebooks", 450),
    ("Spiral Notebook Set", "Diaries & Notebooks", 180),
    ("Bamboo Cover Journal", "Diaries & Notebooks", 320),
    ("Double Wall Ceramic Mug", "Flasks & Mugs", 220),
    ("Stainless Steel Travel Flask", "Flasks & Mugs", 480),
    ("Holi Colour Gift Set", "Holi Specials", 350),
    ("Holi Sweets & Colours Hamper", "Holi Specials", 750),
    ("Wireless Charging Pad", "Tech Items", 650),
    ("USB-C Hub 5-in-1", "Tech Items", 1200),
    ("Wireless Earbuds", "Tech Items", 1450),
    ("Premium Cotton T-Shirt", "Apparel", 320),
    ("Fleece Zip Hoodie", "Apparel", 750),
    ("Canvas Tote Bag", "Bags & Backpacks", 180),
    ('Laptop Backpack 15.6"', "Bags & Backpacks", 980),
    ("Tritan Sports Bottle 750ml", "Bottles", 280),
    ("Copper Insulated Bottle 1L", "Bottles", 520),
    ("Scented Soy Candle", "Home & Living", 380),
    ("Bamboo Desk Organiser", "Home & Living", 450),
    ("Custom Metal Pen", "Cool Extras", 85),
    ("Seed Paper Thank You Cards", "Cool Extras", 45),
    ("Magnetic Phone Stand", "Cool Extras", 290),
]

doc = Document()

title = doc.add_heading("MintBox — Product Catalog", level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph("Name · Category · Price (INR ₹)")
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.runs[0].italic = True

doc.add_paragraph(
    "Source: repository seed catalog (scripts/seed.ts). This reflects the "
    "committed seed data, not a live query of the production Payload CMS."
)

table = doc.add_table(rows=1, cols=3)
table.style = "Light Grid Accent 1"

hdr = table.rows[0].cells
hdr[0].text = "Product"
hdr[1].text = "Category"
hdr[2].text = "Price (₹)"
for cell in hdr:
    cell.paragraphs[0].runs[0].font.bold = True

for name, category, price in PRODUCTS:
    row = table.add_row().cells
    row[0].text = name
    row[1].text = category
    row[2].text = str(price)

doc.add_paragraph()
total = doc.add_paragraph(
    f"Total: {len(PRODUCTS)} products across "
    f"{len({c for _, c, _ in PRODUCTS})} categories."
)
total.runs[0].bold = True

out = "MintBox-Products.docx"
doc.save(out)
print(f"Saved {out}")
